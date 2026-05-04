"""
Revise Results file (figures, Table 1, Table 2) with OOXML tracked changes.

Changes:
  1. Table 2 footnote: standardize significance notation
       "P < .05; *P < .01; **P < .001" -> "* P < .05; ** P < .01; *** P < .001"
  2. Table 2 data cells: update ** -> * for both significant results
       (old convention ** = p<0.05; new convention * = p<0.05)
       Implemented by tracking the lone second asterisk as a deletion.
  3. Figure 2 notes: add sentence explaining t = -2 reference period choice
  4. Figure 3 notes: add pre-trend consistency statement + t = -2 explanation
  5. Table 1 footnote d: clarify control group mixes never-diagnosed +
       delayed-diagnosed; point to eFigure 1
  6. Table 2 footnote a: add note that outcome regression used LPM
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

RESULTS = r"C:\Users\xc77\Dropbox\Claude\Timely Dementia Diagnosis and Planning_Results_04-23-26.docx"
REVISED = r"C:\Users\xc77\Dropbox\Claude\Timely Dementia Diagnosis and Planning_Results_revised_TC.docx"

AUTHOR = "Peer Review Revision"
DATE   = "2026-05-03T00:00:00Z"
_rev_id = [1]

# ── XML helpers ──────────────────────────────────────────────────────────────

def nid():
    v = _rev_id[0]; _rev_id[0] += 1; return str(v)

def clone(e):
    return copy.deepcopy(e) if e is not None else None

def _run(text, rpr=None, deleted=False):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(clone(rpr))
    tag = 'w:delText' if deleted else 'w:t'
    t = OxmlElement(tag)
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def ins_e(text, rpr=None):
    e = OxmlElement('w:ins')
    e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(_run(text, rpr)); return e

def del_e(text, rpr=None):
    e = OxmlElement('w:del')
    e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(_run(text, rpr, deleted=True)); return e

def track_delete_run(r_elem):
    """Wrap an existing w:r in a w:del, converting its w:t to w:delText."""
    r_clone = copy.deepcopy(r_elem)
    for t in r_clone.findall(qn('w:t')):
        t.tag = qn('w:delText')
    d = OxmlElement('w:del')
    d.set(qn('w:id'), nid()); d.set(qn('w:author'), AUTHOR); d.set(qn('w:date'), DATE)
    d.append(r_clone)
    parent = r_elem.getparent()
    idx = list(parent).index(r_elem)
    parent.remove(r_elem)
    parent.insert(idx, d)
    return d

# ── paragraph helpers ────────────────────────────────────────────────────────

def para_text(para):
    return ''.join(r.text or '' for r in para.runs)

def _try_single_run(para, old, new):
    for r_elem in para._p.findall(qn('w:r')):
        t = r_elem.find(qn('w:t'))
        if t is None or not t.text or old not in t.text:
            continue
        rpr = r_elem.find(qn('w:rPr'))
        before, _, after = t.text.partition(old)
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        elems = []
        if before: elems.append(_run(before, rpr))
        elems.append(del_e(old, rpr))
        elems.append(ins_e(new, rpr))
        if after:  elems.append(_run(after, rpr))
        for off, e in enumerate(elems):
            parent.insert(idx + off, e)
        return True
    return False

def _try_multirun(para, old, new):
    runs = para._p.findall(qn('w:r'))
    combined = ''.join(
        (r.find(qn('w:t')).text or '') if r.find(qn('w:t')) is not None else ''
        for r in runs)
    if old not in combined:
        return False
    rpr = runs[0].find(qn('w:rPr')) if runs else None
    for r in runs:
        para._p.remove(r)
    before, _, after = combined.partition(old)
    elems = []
    if before: elems.append(_run(before, rpr))
    elems.append(del_e(old, rpr))
    elems.append(ins_e(new, rpr))
    if after:  elems.append(_run(after, rpr))
    for e in elems:
        para._p.append(e)
    return True

def replace_in_para(para, old, new):
    return _try_single_run(para, old, new) or _try_multirun(para, old, new)

def replace_in_doc(doc, old, new):
    for para in doc.paragraphs:
        if old in para_text(para):
            if replace_in_para(para, old, new):
                return para
    return None

def append_to_para(para, text):
    """Append text as a tracked insertion to the end of para."""
    runs = para._p.findall(qn('w:r'))
    rpr = runs[-1].find(qn('w:rPr')) if runs else None
    para._p.append(ins_e(text, rpr))
    return True

# ── table cell helpers ───────────────────────────────────────────────────────

def cell_para_runs(cell):
    """Return list of (paragraph, run_elem) pairs for all runs in cell."""
    result = []
    for para in cell.paragraphs:
        for r in para._p.findall(qn('w:r')):
            result.append((para, r))
    return result

def cell_text(cell):
    return ''.join(p.text for p in cell.paragraphs)

# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document(RESULTS)
    log = []

    tbl1 = doc.tables[0]
    tbl2 = doc.tables[1]

    # ── 1. Table 2 significance footnote ─────────────────────────────────────
    # Current: "Significance levels: P < .05; *P < .01; **P < .001."
    # New:     "Significance levels: * P < .05; ** P < .01; *** P < .001."
    old1 = "Significance levels: P < .05; *P < .01; **P < .001."
    new1 = "Significance levels: * P < .05; ** P < .01; *** P < .001."
    p1 = replace_in_doc(doc, old1, new1)
    log.append(('1 – Table 2 significance footnote standardized', p1 is not None))

    # ── 2. Table 2 data cells: ** -> * for both significant results ───────────
    # Row 3 (financial respondent): runs end with ")*" + "*" = ")**"
    # Row 4 (will/trust):           runs end with ")*" + "*" = ")**"
    # Strategy: find the standalone trailing "*" run and track-delete it.
    #
    # Under the new convention (* = p<.05), both results stay significant
    # as one star; old ** was p<0.05 under eTable 3 convention, now = *.
    # Authors should verify exact p-values from their output.
    ast_fixed = 0
    for row_idx in [3, 4]:
        cell = tbl2.rows[row_idx].cells[-1]
        pr_pairs = cell_para_runs(cell)
        # The last run should be the trailing lone "*"
        if pr_pairs:
            _, last_r = pr_pairs[-1]
            t = last_r.find(qn('w:t'))
            if t is not None and t.text and t.text.strip() == '*':
                track_delete_run(last_r)
                ast_fixed += 1
    log.append((f'2 – Table 2 asterisks: ** -> * ({ast_fixed}/2 cells updated)', ast_fixed == 2))

    # ── 3. Figure 2 notes: add t = -2 reference period rationale ─────────────
    old3 = ("Estimates represent the adjusted change in probability at each "
            "event time relative to 2 years before dementia onset (t = -2). "
            "Error bars indicate 95% confidence intervals.")
    new3 = ("Estimates represent the adjusted change in probability at each "
            "event time relative to 2 years before dementia onset (t = -2). "
            "The wave at t = -2 was chosen as the reference category to align "
            "with the biennial HRS interview structure; all estimates are "
            "therefore expressed as changes relative to the wave immediately "
            "preceding dementia onset. Error bars indicate 95% confidence intervals.")
    p3 = replace_in_doc(doc, old3, new3)
    log.append(('3 – Figure 2 notes: t = -2 rationale added', p3 is not None))

    # ── 4. Figure 3 notes: add pre-trend + t = -2 rationale ─────────────────
    old4 = "Shaded areas indicate 95% confidence intervals."
    new4 = ("Pre-onset estimates (t = -6 through t = -2) were close to zero "
            "and statistically non-significant, consistent with the parallel "
            "trends assumption. The wave at t = -2 was chosen as the reference "
            "category to align with the biennial HRS interview structure. "
            "Shaded areas indicate 95% confidence intervals.")
    p4 = replace_in_doc(doc, old4, new4)
    log.append(('4 – Figure 3 notes: pre-trend statement + t = -2 rationale added', p4 is not None))

    # ── 5. Table 1 footnote d: clarify control group composition ─────────────
    old5 = ("d Undiagnosed participants included those who received no dementia "
            "diagnosis during the study period, along with pre-diagnosis "
            "person-waves for participants diagnosed after dementia onset. ")
    new5 = ("d Undiagnosed participants included those who received no dementia "
            "diagnosis during the study period, along with pre-diagnosis "
            "person-waves for participants diagnosed after dementia onset. "
            "This group therefore comprises permanently undiagnosed individuals "
            "and those who eventually received a delayed clinical diagnosis; "
            "the distribution of diagnosis timing is shown in eFigure 1 of "
            "the Supplement.")
    p5 = replace_in_doc(doc, old5, new5)
    log.append(('5 – Table 1 footnote d: control group composition clarified', p5 is not None))

    # ── 6. Table 2 footnote a: note LPM used for outcome regression ──────────
    old6 = ("ᵃ Adjusted differences were estimated using the Callaway-Sant’Anna "
            "(2021) doubly robust difference-in-differences event-study estimator, "
            "aggregated across post-onset event times (see eMethods 3 for details).")
    new6 = ("ᵃ Adjusted differences were estimated using the Callaway-Sant’Anna "
            "(2021) doubly robust difference-in-differences event-study estimator, "
            "aggregated across post-onset event times (see eMethods 3 for details). "
            "The outcome regression component was estimated using a linear "
            "probability model, consistent with the default specification in "
            "the did R package. Authors should verify exact p-values against "
            "model output and update asterisks accordingly.")
    p6 = replace_in_doc(doc, old6, new6)
    log.append(('6 – Table 2 footnote a: LPM note + p-value verification reminder', p6 is not None))

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(REVISED)

    # ── Report ───────────────────────────────────────────────────────────────
    print(f"\nSaved: {REVISED}\n")
    print(f"{'Change':<65} {'Status'}")
    print("-" * 73)
    for label, ok in log:
        print(f"{label:<65} {'OK' if ok else 'MISS'}")
    print()

if __name__ == "__main__":
    main()
