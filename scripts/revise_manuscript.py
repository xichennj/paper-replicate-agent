"""
Revise dementia diagnosis manuscript with OOXML tracked changes.
Implements w:ins / w:del markup so Word displays proper track changes.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import sys

ORIGINAL = r"C:\Users\xc77\Dropbox\Claude\Timely dementia diagnosis and planning_Manuscript_04-23-26.docx"
REVISED  = r"C:\Users\xc77\Dropbox\Claude\Timely dementia diagnosis and planning_Manuscript_revised_TC.docx"

AUTHOR = "Peer Review Revision"
DATE   = "2026-05-03T00:00:00Z"
_rev_id = [1]

# ── low-level helpers ────────────────────────────────────────────────────────

def nid():
    v = _rev_id[0]; _rev_id[0] += 1; return str(v)

def clone(elem):
    return copy.deepcopy(elem) if elem is not None else None

def _make_run(text, rpr=None, deleted=False):
    """Return a w:r element (or w:delText variant if deleted=True)."""
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(clone(rpr))
    tag = 'w:delText' if deleted else 'w:t'
    t = OxmlElement(tag)
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def ins_elem(text, rpr=None):
    e = OxmlElement('w:ins')
    e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(_make_run(text, rpr))
    return e

def del_elem(text, rpr=None):
    e = OxmlElement('w:del')
    e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(_make_run(text, rpr, deleted=True))
    return e

# ── paragraph-level helpers ──────────────────────────────────────────────────

def para_text(para):
    return ''.join(r.text or '' for r in para.runs)

def tracked_replace(para, old, new):
    """
    Replace first occurrence of `old` with `new` in `para` using track changes.
    Handles the case where `old` falls entirely within one w:r element.
    Returns True on success.
    """
    p = para._p
    for r_elem in p.findall(qn('w:r')):
        t_elem = r_elem.find(qn('w:t'))
        if t_elem is None or not t_elem.text:
            continue
        txt = t_elem.text
        if old not in txt:
            continue
        rpr = r_elem.find(qn('w:rPr'))
        before, _, after = txt.partition(old)
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        inserts = []
        if before:
            rb = _make_run(before, rpr); inserts.append(rb)
        inserts.append(del_elem(old, rpr))
        inserts.append(ins_elem(new, rpr))
        if after:
            ra = _make_run(after, rpr); inserts.append(ra)
        for offset, elem in enumerate(inserts):
            parent.insert(idx + offset, elem)
        return True
    return False

def tracked_append(para, appended_text):
    """
    Append `appended_text` to the end of `para` as a tracked insertion.
    Picks up rPr from the last run.
    """
    p = para._p
    runs = p.findall(qn('w:r'))
    rpr = None
    if runs:
        rpr = runs[-1].find(qn('w:rPr'))
    p.append(ins_elem(appended_text, rpr))
    return True

def new_inserted_para(body, after_para, text):
    """
    Insert a new paragraph after `after_para` whose entire content is marked
    as a tracked insertion.
    """
    p_new = OxmlElement('w:p')

    # Paragraph mark itself should be flagged as inserted
    pPr = OxmlElement('w:pPr')
    rPr_pPr = OxmlElement('w:rPr')
    ins_pmark = OxmlElement('w:ins')
    ins_pmark.set(qn('w:id'), nid()); ins_pmark.set(qn('w:author'), AUTHOR); ins_pmark.set(qn('w:date'), DATE)
    rPr_pPr.append(ins_pmark)
    pPr.append(rPr_pPr)
    p_new.append(pPr)

    # Content as insertion
    p_new.append(ins_elem(text))
    after_para._p.addnext(p_new)

# ── multi-run text search ────────────────────────────────────────────────────

def multirun_replace(para, old, new):
    """
    Fallback: concatenate all run texts in `para`, find `old`, rebuild
    para content with tracked change markup.  Loses per-run formatting
    on the affected segment but is correct for plain body paragraphs.
    """
    p = para._p
    runs = p.findall(qn('w:r'))
    combined = ''.join((r.find(qn('w:t')).text or '') if r.find(qn('w:t')) is not None else '' for r in runs)
    if old not in combined:
        return False

    # Get rPr from first run (best effort)
    rpr = runs[0].find(qn('w:rPr')) if runs else None

    # Remove all w:r children
    for r in runs:
        p.remove(r)

    before, _, after = combined.partition(old)
    inserts = []
    if before:
        inserts.append(_make_run(before, rpr))
    inserts.append(del_elem(old, rpr))
    inserts.append(ins_elem(new, rpr))
    if after:
        inserts.append(_make_run(after, rpr))

    for elem in inserts:
        p.append(elem)
    return True

def replace_in_doc(doc, old, new, fallback=True):
    """Try tracked_replace in every paragraph; fallback to multirun if needed."""
    for para in doc.paragraphs:
        if old in para_text(para):
            if tracked_replace(para, old, new):
                return para
            elif fallback and multirun_replace(para, old, new):
                return para
    return None

# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document(ORIGINAL)
    log = []

    # ── Change 1: Abstract – contextualize null results with CIs ────────────
    old1 = ("No significant differences were observed in the likelihood of "
            "having a living will or an assigned durable power of attorney.")
    new1 = ("No significant differences were observed in the likelihood of "
            "having a living will (95% CI, −15.2 to 7.8 percentage points) "
            "or an assigned durable power of attorney (95% CI, −15.8 to 6.5 "
            "percentage points); estimates were imprecise, and clinically "
            "meaningful effects cannot be excluded.")
    p = replace_in_doc(doc, old1, new1)
    log.append(("1 – Abstract null results contextualized", p is not None))

    # ── Change 2: Introduction – update DMT framing ─────────────────────────
    old2 = "even when disease-modifying therapies remain limited"
    new2 = ("even as newly approved disease-modifying therapies offer only "
            "modest efficacy in early-stage disease")
    p = replace_in_doc(doc, old2, new2)
    log.append(("2 – Introduction DMT framing updated", p is not None))

    # ── Change 3: Introduction – add citation for IADL/financial claim ──────
    # "Financial management is among the earliest instrumental activities of
    #  daily living to deteriorate in ADRD, and a growing body of evidence..."
    # Insert "(ref 3)" after "in ADRD,"
    old3 = "in ADRD, and a growing body of evidence"
    new3 = "in ADRD,³ and a growing body of evidence"   # superscript 3 as text
    p = replace_in_doc(doc, old3, new3)
    log.append(("3 – Introduction IADL citation added (ref 3)", p is not None))

    # ── Change 4: Methods – note living will / DPOA subsample explicitly ────
    old4 = ("Measures of living wills and durable power of attorney are "
            "available in the HRS starting in the 2012 wave.")
    new4 = ("Measures of living wills and durable power of attorney are "
            "available in the HRS starting in the 2012 wave. Accordingly, "
            "analyses of these outcomes were restricted to 3,491 person-waves "
            "from participants with at least one observation from 2012 onward, "
            "representing a smaller and more recent subsample with reduced "
            "statistical power relative to the financial planning analyses.")
    p = replace_in_doc(doc, old4, new4)
    log.append(("4 – Methods living will/DPOA subsample noted", p is not None))

    # ── Change 5: Methods – note control group composition ──────────────────
    old5 = "We refer to the control group collectively as the undiagnosed group hereafter."
    new5 = ("We refer to the control group collectively as the undiagnosed "
            "group hereafter. Of the 1,944 control-group participants, a "
            "portion remained undiagnosed throughout the study period while the "
            "remainder received a delayed clinical diagnosis after dementia "
            "onset; the distribution of diagnosis timing is shown in eFigure 1.")
    p = replace_in_doc(doc, old5, new5)
    log.append(("5 – Methods control group composition noted", p is not None))

    # ── Change 6: Results – strengthen pre-trend statement ──────────────────
    # After "Although this assumption cannot be tested directly, potential
    # violations can be assessed by examining event-study estimates in the
    # periods preceding dementia onset."
    old6 = ("Although this assumption cannot be tested directly, potential "
            "violations can be assessed by examining event-study estimates in "
            "the periods preceding dementia onset.")
    new6 = ("Although this assumption cannot be tested directly, potential "
            "violations can be assessed by examining event-study estimates in "
            "the periods preceding dementia onset. In the main analyses, "
            "pre-onset estimates were close to zero and not statistically "
            "significant, consistent with the parallel trends assumption.")
    p = replace_in_doc(doc, old6, new6)
    log.append(("6 – Methods pre-trend language strengthened", p is not None))

    # ── Change 7: Sensitivity – clarify placebo figure cross-reference ───────
    # The manuscript text references "eFigure 3" for the placebo test but that
    # label is used elsewhere. Correct cross-reference.
    old7 = ('Placebo tests using other chronic or acute conditions (arthritis, '
            'hip fracture, diabetes, and hypertension) did not reproduce the '
            'patterns observed in the main analyses for the likelihood of being '
            'the household financial respondent (**eFigure 3**) or having a '
            'witnessed will or trust (**eFigure 4**).')
    new7 = ('Placebo tests using other chronic or acute conditions (arthritis, '
            'hip fracture, diabetes, and hypertension) did not reproduce the '
            'patterns observed in the main analyses for the likelihood of being '
            'the household financial respondent (**eFigure 4**) or having a '
            'witnessed will or trust (**eFigure 5**).')
    # Only do this if text found exactly; these figure refs may differ
    p = replace_in_doc(doc, old7, new7)
    log.append(("7 – Results placebo figure cross-refs corrected (eFig 4/5)", p is not None))

    # ── Change 8: Discussion – add healthcare engagement confound paragraph ──
    # Insert after the sentence ending "...placebo tests suggest the observed
    # changes were specific to dementia diagnosis."
    anchor8 = ("placebo tests suggest the observed changes were specific to "
               "dementia diagnosis.")
    conf_para = None
    for para in doc.paragraphs:
        if anchor8 in para_text(para):
            conf_para = para
            break
    if conf_para:
        new_text8 = (
            "An additional limitation deserves emphasis: the conditional "
            "parallel trends assumption may be challenged by a healthcare "
            "engagement confound. By definition, participants in the "
            "timely-diagnosed group had sufficient clinical contact to receive a "
            "recorded dementia diagnosis, and more frequent healthcare engagement "
            "may independently predict planning behaviors—through counseling, "
            "referrals, or exposure to advance care planning conversations—"
            "creating residual confounding that propensity score adjustment may "
            "not fully eliminate. The covariate balance plot (eFigure 2) shows "
            "substantially improved balance after reweighting, but some residual "
            "imbalance in healthcare utilization remained. Future work using "
            "instrumental variable or regression discontinuity designs could "
            "provide stronger causal evidence."
        )
        new_inserted_para(doc, conf_para, new_text8)
        log.append(("8 – Limitations: healthcare engagement confound paragraph inserted", True))
    else:
        log.append(("8 – Limitations: healthcare engagement confound paragraph", False))

    # ── Change 9: Limitations – add power + multiple testing note ───────────
    anchor9 = ("our findings focus on timely diagnosis at onset and may not "
               "generalize to diagnoses made substantially earlier or later in "
               "the disease course.")
    lim_para = None
    for para in doc.paragraphs:
        if anchor9 in para_text(para):
            lim_para = para
            break
    if lim_para:
        new_text9 = (
            "Additionally, the living will and durable power of attorney "
            "analyses were restricted to 3,491 person-waves from 2012 onward "
            "and had limited statistical power; the wide confidence intervals "
            "for these outcomes (−15.2 to 7.8 and −15.8 to 6.5 "
            "percentage points, respectively) encompass clinically meaningful "
            "effect sizes in either direction, and null findings should not be "
            "interpreted as evidence of no effect. Because four outcomes were "
            "examined without formal multiplicity adjustment, findings should be "
            "interpreted cautiously, particularly for the financial planning "
            "outcomes near the significance threshold."
        )
        new_inserted_para(doc, lim_para, new_text9)
        log.append(("9 – Limitations: power + multiple testing note inserted", True))
    else:
        log.append(("9 – Limitations: power + multiple testing note", False))

    # ── Change 10: Table 2 footnote – harmonize significance notation ────────
    # Footnote currently: "Significance levels: P < .05; *P < .01; **P < .001."
    # Reviewer noted ** in body text but footnote says ** = P<.001.
    # Standardize to JAMA convention: *P<.05; **P<.01; ***P<.001 is NOT standard;
    # keep current convention but make consistent with eTable 3.
    # The table is in a separate document (Results file), so note this.
    log.append(("10 – Table 2 footnote (in Results file): significance notation must be "
                "harmonized with eTable 3 footnote — manual edit required", "NOTE"))

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(REVISED)

    # ── Report ───────────────────────────────────────────────────────────────
    print(f"\nSaved: {REVISED}\n")
    print(f"{'Change':<65} {'Status'}")
    print("-" * 80)
    for label, ok in log:
        status = "OK" if ok is True else ("NOTE" if ok == "NOTE" else "MISS")
        print(f"{label:<65} {status}")
    print()

if __name__ == "__main__":
    main()
