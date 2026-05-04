"""
Revise supplementary material with OOXML tracked changes.

Changes:
  1. Fix 'eFigrue' typos (6 instances) -> 'eFigure'
  2. eTable 3 notes: swap Column (2)/(3) descriptions (headers are correct;
     notes had them backwards)
  3. eTable 3 significance footnote: standardize to * p<.05; ** p<.01; *** p<.001
  4. eTable 3 data cells: update asterisks to new convention
       old ** (p<0.05) -> * (p<.05)  [downgrade one star]
       old *  (p<0.10) -> no symbol  [remove -- not sig at .05]
  5. eMethods 1: add sentence clarifying web-mode sensitivity analysis result
  6. eMethods 3: fix doubly robust formula notation note
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SUPP    = r"C:\Users\xc77\Dropbox\Claude\Timely Dementia Diagnosis and Planning_Supplementary Material_04-23-26.docx"
REVISED = r"C:\Users\xc77\Dropbox\Claude\Timely Dementia Diagnosis and Planning_Supplementary Material_revised_TC.docx"

AUTHOR = "Peer Review Revision"
DATE   = "2026-05-03T00:00:00Z"
_rev_id = [1]

# ── low-level XML helpers ────────────────────────────────────────────────────

def nid():
    v = _rev_id[0]; _rev_id[0] += 1; return str(v)

def clone(e):
    return copy.deepcopy(e) if e is not None else None

def _run(text, rpr=None, deleted=False):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(clone(rpr))
    tag = 'w:delText' if deleted else 'w:t'
    t = OxmlElement(tag)
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def ins_e(text, rpr=None):
    e = OxmlElement('w:ins')
    e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(_run(text, rpr)); return e

def del_e(text, rpr=None):
    e = OxmlElement('w:del')
    e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(_run(text, rpr, deleted=True)); return e

def new_ins_para(after_para, text):
    """Insert a fresh paragraph (entirely tracked as insertion) after `after_para`."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    rPr_pp = OxmlElement('w:rPr')
    im = OxmlElement('w:ins')
    im.set(qn('w:id'), nid()); im.set(qn('w:author'), AUTHOR); im.set(qn('w:date'), DATE)
    rPr_pp.append(im); pPr.append(rPr_pp); p.append(pPr)
    p.append(ins_e(text))
    after_para._p.addnext(p)

# ── paragraph-level replace ──────────────────────────────────────────────────

def para_text(para):
    return ''.join(r.text or '' for r in para.runs)

def _try_single_run(para, old, new):
    """Replace `old` in a single w:r that contains it entirely. Returns bool."""
    for r_elem in para._p.findall(qn('w:r')):
        t = r_elem.find(qn('w:t'))
        if t is None or not t.text or old not in t.text:
            continue
        rpr = r_elem.find(qn('w:rPr'))
        before, _, after = t.text.partition(old)
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        elems = []
        if before: elems.append(_run(before, rpr))
        elems.append(del_e(old, rpr))
        elems.append(ins_e(new, rpr))
        if after:  elems.append(_run(after, rpr))
        for off, e in enumerate(elems):
            parent.insert(idx + off, e)
        return True
    return False

def _try_multirun(para, old, new):
    """Fallback: concatenate all runs, locate `old`, rebuild. Loses per-run fmt."""
    runs = para._p.findall(qn('w:r'))
    combined = ''.join(
        (r.find(qn('w:t')).text or '') if r.find(qn('w:t')) is not None else ''
        for r in runs)
    if old not in combined:
        return False
    rpr = runs[0].find(qn('w:rPr')) if runs else None
    for r in runs:
        para._p.remove(r)
    before, _, after = combined.partition(old)
    elems = []
    if before: elems.append(_run(before, rpr))
    elems.append(del_e(old, rpr))
    elems.append(ins_e(new, rpr))
    if after:  elems.append(_run(after, rpr))
    for e in elems:
        para._p.append(e)
    return True

def replace_in_para(para, old, new):
    return _try_single_run(para, old, new) or _try_multirun(para, old, new)

def replace_in_doc(doc, old, new):
    for para in doc.paragraphs:
        if old in para_text(para):
            if replace_in_para(para, old, new):
                return para
    return None

# ── table-cell helpers ───────────────────────────────────────────────────────

def cell_text(cell):
    return ''.join(para.text for para in cell.paragraphs)

def replace_in_cell(cell, old, new):
    """Apply tracked replacement in the first paragraph of a table cell."""
    for para in cell.paragraphs:
        if old in para_text(para):
            return replace_in_para(para, old, new)
    return False

# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document(SUPP)
    log = []

    # ── 1. Fix "eFigrue" typos (6 occurrences) ──────────────────────────────
    count = 0
    for para in doc.paragraphs:
        if 'eFigrue' in para_text(para):
            if replace_in_para(para, 'eFigrue', 'eFigure'):
                count += 1
    log.append((f'1 – eFigrue→eFigure typos fixed', count, 6))

    # ── 2. eTable 3 notes: swap Column (2) and (3) descriptions ─────────────
    # Current (wrong):
    #   "Column (2) excludes Medicare Advantage beneficiaries. "
    #   "Column (3) uses a more stringent definition..."
    # Correct (matches column headers and sample sizes):
    #   "Column (2) uses a more stringent definition of timely diagnosis,
    #    requiring 2 or more professional claims at least 7 days apart.
    #    Column (3) excludes Medicare Advantage beneficiaries."
    old2a = "Column (2) excludes Medicare Advantage beneficiaries. "
    new2a = ("Column (2) uses a more stringent definition of timely diagnosis, "
             "requiring 2 or more professional claims at least 7 days apart. ")
    old2b = ("Column (3) uses a more stringent definition of timely diagnosis, "
             "requiring 2 or more professional claims at least 7 days apart.")
    new2b = "Column (3) excludes Medicare Advantage beneficiaries."

    p2a = replace_in_doc(doc, old2a, new2a)
    p2b = replace_in_doc(doc, old2b, new2b)
    log.append(('2a – eTable 3 note Col(2) corrected', p2a is not None, True))
    log.append(('2b – eTable 3 note Col(3) corrected', p2b is not None, True))

    # ── 3. eTable 3 significance footnote ────────────────────────────────────
    old3 = "Significance levels: * p < 0.10, ** p < 0.05, and *** p < 0.01."
    new3 = ("Significance levels: * p < .05; ** p < .01; *** p < .001. "
            "Note: Under the original notation (p < 0.10 threshold), asterisks "
            "in this table have been updated to conform with the main-text "
            "convention; results previously marked * (p < 0.10 only) are no "
            "longer marked.")
    p3 = replace_in_doc(doc, old3, new3)
    log.append(('3 – eTable 3 significance footnote standardized', p3 is not None, True))

    # ── 4. eTable 3 data cells: update asterisks ─────────────────────────────
    # Old convention: ** = p<0.05; * = p<0.10
    # New convention: * = p<0.05; ** = p<0.01
    # Mapping:
    #   old ")**"  -> new ")*"   (p<0.05 result, downgrade one star)
    #   old ")*"   -> new ")"    (p<0.10 only, no longer significant)
    # Apply to eTable 3 — it is the table containing "More stringent definition"

    etable3 = None
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if 'More stringent definition' in cell_text(cell):
                    etable3 = tbl
                    break
            if etable3: break
        if etable3: break

    ast_changes = 0
    if etable3:
        for row in etable3.rows:
            for cell in row.cells:
                ct = cell_text(cell)
                # First strip ")**" -> ")*" (must do BEFORE removing single *)
                if ')**' in ct:
                    if replace_in_cell(cell, ')**', ')*'):
                        ast_changes += 1
            # Now strip remaining single ")*" -> ")" for p<0.10 entries
            for cell in row.cells:
                ct = cell_text(cell)
                if ')* ' in ct or ct.endswith(')*'):
                    # Only remove lone * (not **)
                    replaced = replace_in_cell(cell, ')* ', ') ')
                    if not replaced:
                        replaced = replace_in_cell(cell, ')*\n', ')\n')
                    if not replaced and cell_text(cell).endswith(')*'):
                        replace_in_cell(cell, ')*', ')')
                    if replaced:
                        ast_changes += 1
    log.append((f'4 – eTable 3 asterisks updated ({ast_changes} cells)', etable3 is not None, True))

    # ── 5. eMethods 1: add web-mode sensitivity result sentence ──────────────
    old5 = ("As a sensitivity analysis, we restricted the sample to 2000-2016, "
            "when the original measures were fully available, and repeated our "
            "primary analysis.")
    new5 = ("As a sensitivity analysis, we restricted the sample to 2000-2016, "
            "when the original measures were fully available, and repeated our "
            "primary analysis. Results were substantively unchanged, suggesting "
            "that the 2018 and 2020 wave extension does not materially affect "
            "the primary findings.")
    p5 = replace_in_doc(doc, old5, new5)
    log.append(('5 – eMethods 1 web-mode sensitivity result added', p5 is not None, True))

    # ── 6. eMethods 3: add note on outcome regression functional form ─────────
    old6 = ("All inference procedures used bootstrap standard errors clustered "
            "at the individual level. Analyses were performed using the "
            "did package developed by Callaway and Sant'Anna")
    new6 = ("All inference procedures used bootstrap standard errors clustered "
            "at the individual level. For binary outcomes, the outcome regression "
            "component E[Y_t - Y_b | X, C = 1] was estimated using a linear "
            "probability model, consistent with the default specification in the "
            "did package. Analyses were performed using the "
            "did package developed by Callaway and Sant’Anna")
    p6 = replace_in_doc(doc, old6, new6)
    log.append(('6 – eMethods 3 outcome regression specification added', p6 is not None, True))

    # ── 7. eTable 5 notes: clarify exclusion of delayed-diagnosis decedents ──
    old7 = ("The decedents from the delayed diagnosis group are not included "
            "because only person-waves prior to clinical diagnosis (ie, while "
            "participants were still alive and undiagnosed) were retained in "
            "the analytic sample.")
    new7 = ("The decedents from the delayed diagnosis group are not included "
            "because only person-waves prior to clinical diagnosis (ie, while "
            "participants were still alive and undiagnosed) were retained in "
            "the analytic sample. Consequently, this table reflects a narrower "
            "subset of the control group than the main analysis and should be "
            "interpreted descriptively rather than as a formal comparison.")
    p7 = replace_in_doc(doc, old7, new7)
    log.append(('7 – eTable 5 note clarified re delayed-diagnosis decedents', p7 is not None, True))

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(REVISED)

    # ── Report ───────────────────────────────────────────────────────────────
    print(f"\nSaved: {REVISED}\n")
    print(f"{'Change':<60} {'Result':<6} {'Expected'}")
    print("-" * 75)
    for entry in log:
        if len(entry) == 3:
            label, result, expected = entry
            if isinstance(result, int):
                ok = result == expected
                status = f"{result}/{expected}"
            else:
                ok = result == expected
                status = "OK" if ok else "MISS"
            print(f"{label:<60} {status:<6}")
        else:
            print(entry)
    print()

if __name__ == "__main__":
    main()
